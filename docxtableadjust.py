# run: python docxadjust.py testtable.docx --auto-detect

import docx
from docx.shared import Inches, Mm, Pt
import sys
import os
from docx.enum.table import WD_TABLE_ALIGNMENT

def get_a4_dimensions(margin_mm=25.4):
    """
    Returns the dimensions of an A4 page in inches, accounting for margins.
    Standard A4 is 210mm × 297mm.
    Default margins: 1 inch (25.4mm) on all sides.
    
    Args:
        margin_mm: Margin size in millimeters (default: 25.4mm = 1 inch)
    """
    a4_width_mm = 210  # mm
    a4_height_mm = 297  # mm
    
    available_width_mm = a4_width_mm - (2 * margin_mm)
    available_height_mm = a4_height_mm - (2 * margin_mm)
    
    return Mm(available_width_mm), Mm(available_height_mm)

def handle_complex_table_structure(table, max_width):
    """
    Handle tables with variable column counts across rows or nested structure.
    
    Args:
        table: The docx table object
        max_width: Maximum available width
        
    Returns:
        The adjusted table
    """
    if not table.rows or len(table.rows) == 0:
        return table
    
    # Find the maximum number of columns in any row
    max_cols = 0
    for row in table.rows:
        max_cols = max(max_cols, len(row.cells))
    
    if max_cols == 0:
        return table
    
    # Identify special row patterns to determine custom width distributions
    # For program spec tables like in the example
    
    # Pattern 1: Program spec table with ID and Mode (3 columns in first row)
    if max_cols >= 3:
        # Check if first row has a 3-column structure with 'Program ID' and 'Mode'
        has_program_id_mode = False
        for i, row in enumerate(table.rows):
            if len(row.cells) >= 3:
                # Check for 'Program ID' text in first cell
                if 'program id' in row.cells[0].text.lower():
                    # And 'Mode' in the third cell
                    if len(row.cells) >= 3 and 'mode' in row.cells[2].text.lower():
                        has_program_id_mode = True
                        break
        
        if has_program_id_mode:
            # This is a program spec table with ID and Mode format
            # Apply special column distribution: wide first column, medium middle column, narrow mode column
            col_ratios = [3, 4, 1]
            
            # Calculate column widths
            total_ratio = sum(col_ratios[:max_cols])
            col_widths = [max_width * (ratio / total_ratio) for ratio in col_ratios[:max_cols]]
            
            # Apply widths for each row, respecting the number of cells in each row
            for row in table.rows:
                cells_in_row = len(row.cells)
                
                if cells_in_row == 1:
                    # Full-width row (section header or blank row)
                    row.cells[0].width = max_width
                elif cells_in_row == 2:
                    # Two-column row (typical field-value pair)
                    # Use widths from first two columns in the pattern
                    combined_last_cols = sum(col_widths[1:])
                    row.cells[0].width = col_widths[0]
                    row.cells[1].width = combined_last_cols
                elif cells_in_row >= 3:
                    # Apply the 3-column distribution
                    for i in range(min(cells_in_row, len(col_widths))):
                        row.cells[i].width = col_widths[i]
            
            return table
    
    # Pattern 2: Amendment History table (typically 4 columns with specific headers)
    has_amendment_history = False
    for i, row in enumerate(table.rows):
        if i < len(table.rows) - 1 and len(row.cells) >= 1:
            if 'amendment history' in row.cells[0].text.lower():
                has_amendment_history = True
                break
    
    if has_amendment_history and max_cols >= 4:
        # For amendment history tables, use a specific column distribution
        # Change Number | Revision Description | Revision Number | Date
        col_ratios = [1.5, 3, 1.5, 1]
        
        # Make sure the ratios match the actual max columns
        while len(col_ratios) > max_cols:
            col_ratios.pop()
        
        # Calculate column widths
        total_ratio = sum(col_ratios)
        col_widths = [max_width * (ratio / total_ratio) for ratio in col_ratios]
        
        # Apply to all rows with max_cols
        for row in table.rows:
            if len(row.cells) == max_cols:
                for i, cell in enumerate(row.cells):
                    cell.width = col_widths[i]
            elif len(row.cells) == 1:
                # Full-width row (section header)
                row.cells[0].width = max_width
        
        return table
    
    # Default handling for other table structures
    # For tables with variable column counts, assign proportional widths
    # based on the maximum column count
    default_col_width = max_width / max_cols
    
    # Start with equal distribution
    for row in table.rows:
        cells_in_row = len(row.cells)
        
        if cells_in_row == 1:
            # Full-width row
            row.cells[0].width = max_width
        elif cells_in_row == 2 and max_cols > 2:
            # If this is a two-column row in a table with more columns,
            # assume it's a label-value pair and make first column narrower
            row.cells[0].width = max_width * 0.3
            row.cells[1].width = max_width * 0.7
        else:
            # Distribute width evenly among cells
            for cell in row.cells:
                cell.width = default_col_width
    
    return table

def fix_pandoc_tables(doc):
    """
    Apply specific fixes for pandoc-generated tables which often have structure issues.
    
    Args:
        doc: The docx document object
    """
    # Find all tables in the document
    tables_to_fix = []
    
    # Collect tables from main document
    for table in doc.tables:
        tables_to_fix.append(table)
    
    # Also look for tables in headers, footers, and sections
    for section in doc.sections:
        try:
            # Check header
            for table in section.header.tables:
                tables_to_fix.append(table)
            
            # Check footer
            for table in section.footer.tables:
                tables_to_fix.append(table)
        except Exception:
            # Some sections might not have headers/footers
            pass
    
    # Apply fixes to each table
    for table in tables_to_fix:
        try:
            # Ensure table structure is sound
            if hasattr(table, '_tbl'):
                # Check for and fix missing or corrupted properties
                if not hasattr(table._tbl, 'tblPr'):
                    # This is a common issue with pandoc tables
                    from docx.oxml.table import CT_Tbl, CT_TblPr
                    if isinstance(table._tbl, CT_Tbl):
                        if table._tbl.tblPr is None:
                            table._tbl.tblPr = CT_TblPr()
        except Exception as e:
            print(f"Warning: Could not fix table structure: {e}")
    
    return doc

def adjust_table_size(table, max_width, column_widths=None):
    """
    Adjust the table width to fit within the provided max_width.
    
    Args:
        table: The docx table object to adjust
        max_width: Maximum available width in docx units
        column_widths: Optional list of relative width values for each column.
                      If provided, must have same length as the number of columns.
                      Example: [1, 2, 1] would make the middle column twice as wide as the others.
                      Can also be a string like 'complex-program-spec' to use special handling.
    """
    # Safety check - ensure table has rows and columns
    if not table.rows or len(table.rows) == 0:
        print("Warning: Table has no rows, skipping adjustment")
        return table
    
    # Check for complex table marker
    if isinstance(column_widths, str) and column_widths == 'complex-program-spec':
        return handle_complex_table_structure(table, max_width)
        
    # Count the number of columns (check all rows to ensure consistency)
    num_cols = 0
    for row in table.rows:
        num_cols = max(num_cols, len(row.cells))
    
    if num_cols == 0:
        print("Warning: Table has no columns, skipping adjustment")
        return table
    
    # Check for variable column structure
    has_variable_columns = False
    for row in table.rows:
        if len(row.cells) > 0 and len(row.cells) != num_cols:
            has_variable_columns = True
            break
    
    # If variable column structure detected, use complex handling
    if has_variable_columns:
        return handle_complex_table_structure(table, max_width)
    
    # For regular tables, proceed with normal handling
    
    # Attempt to set table width directly if possible
    try:
        # Set table to exactly max width
        table.width = max_width
    except AttributeError:
        # Some older versions of python-docx don't have direct table width setting
        # We'll need to adjust column widths individually
        pass
    
    # Ensure table is centered
    try:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    except AttributeError:
        # Handle if alignment property isn't available
        pass
    
    # Fix table grid to ensure consistent structure
    try:
        # This can help with pandoc-generated tables
        if hasattr(table, '_tbl') and hasattr(table._tbl, 'tblGrid'):
            # Ensure grid has correct number of columns
            grid = table._tbl.tblGrid
            while len(grid.gridCol_lst) < num_cols:
                grid.add_gridCol()
            while len(grid.gridCol_lst) > num_cols:
                grid.remove(grid.gridCol_lst[-1])
    except Exception as e:
        print(f"Warning: Could not fix table grid: {e}")
    
    # Process column widths
    if column_widths and len(column_widths) == num_cols:
        # Calculate the sum of all relative widths
        total_relative_width = sum(column_widths)
        
        # Calculate the absolute width for each column based on relative proportions
        absolute_widths = [max_width * (rw / total_relative_width) for rw in column_widths]
        
        # Apply the calculated widths to each column
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i < len(absolute_widths):
                    try:
                        cell.width = absolute_widths[i]
                    except Exception as e:
                        print(f"Warning: Could not set cell width: {e}")
    else:
        # If no specific widths provided, distribute evenly
        col_width = max_width / num_cols
        
        # Set each column to equal width for all rows
        for row in table.rows:
            for cell in row.cells:
                try:
                    cell.width = col_width
                except Exception as e:
                    print(f"Warning: Could not set cell width: {e}")
    
    # Set auto-fit to False to maintain our specific widths
    try:
        table.autofit = False
    except AttributeError:
        # Some versions don't have this attribute
        pass
    
    # Additional fixes specific to pandoc-converted tables
    try:
        # Set table layout to fixed instead of auto
        for row in table.rows:
            for cell in row.cells:
                # Ensure text wraps within cells
                for paragraph in cell.paragraphs:
                    paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.LEFT
                    # Set paragraph spacing to single (1.0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    
                    if not paragraph.runs:
                        continue
                    for run in paragraph.runs:
                        # Set Times New Roman font and size 12
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
    except Exception as e:
        print(f"Warning: Could not apply additional fixes: {e}")
    
    return table

def detect_table_patterns(doc):
    """
    Analyze the document and detect common table patterns to suggest column widths.
    Especially useful for pandoc-generated markdown tables.
    
    Args:
        doc: The docx document object
        
    Returns:
        dict: Dictionary mapping table indices to suggested column widths
    """
    suggested_widths = {}
    
    for i, table in enumerate(doc.tables):
        # Skip empty tables
        if not table.rows or len(table.rows) == 0:
            continue
            
        # Find the max number of columns in any row
        max_cols = 0
        for row in table.rows:
            max_cols = max(max_cols, len(row.cells))
        
        if max_cols == 0:
            continue
        
        # Check for variable column structure (suggesting a complex table)
        has_variable_columns = False
        for row in table.rows:
            if len(row.cells) > 0 and len(row.cells) != max_cols:
                has_variable_columns = True
                break
        
        # Check if this table has a 'Program ID' or 'Mode' cell
        has_program_id = False
        has_mode = False
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip().lower()
                if 'program id' in text:
                    has_program_id = True
                if 'mode' in text:
                    has_mode = True
        
        # Check for amendment history section
        has_amendment_history = False
        for row in table.rows:
            if len(row.cells) > 0:
                text = row.cells[0].text.strip().lower()
                if 'amendment history' in text:
                    has_amendment_history = True
                    break
        
        # Suggest appropriate column widths based on detected patterns
        if has_variable_columns and has_program_id:
            # This is likely a program specification table with variable structure
            suggested_widths[i] = 'complex-program-spec'
            continue
        
        if has_amendment_history and max_cols >= 3:
            # This is likely an amendment history table
            if max_cols == 3:
                suggested_widths[i] = [1.5, 3, 1.5]
            elif max_cols >= 4:
                suggested_widths[i] = [1.5, 3, 1.5, 1]
            continue
            
        # Check if this looks like a program spec table (from the image examples)
        # These often have a narrow first column and a wide second column
        if max_cols == 2 and not has_variable_columns:
            # Check cell contents to detect if it's a spec table
            is_spec_table = False
            try:
                # Look for common program spec fields in first column cells
                first_col_texts = []
                for row in table.rows:
                    if len(row.cells) >= 1:
                        first_col_texts.append(row.cells[0].text.strip().lower())
                
                spec_keywords = ['program', 'name', 'description', 'environment', 'source', 
                                'language', 'compiler', 'id']
                
                # Count how many spec keywords are in the first column
                keyword_matches = sum(1 for text in first_col_texts if any(kw in text for kw in spec_keywords))
                
                # If more than 2 matches, likely a spec table
                if keyword_matches >= 2:
                    is_spec_table = True
                    # For spec tables, make the second column wider (1:3 ratio)
                    suggested_widths[i] = [1, 3]
            except Exception:
                pass
                
        # Check if this looks like a validation table (3 columns with headers like Input, Validation, Remarks)
        elif max_cols == 3 and not has_variable_columns:
            try:
                if len(table.rows) > 0:
                    header_texts = []
                    for cell in table.rows[0].cells:
                        header_texts.append(cell.text.strip().lower())
                    
                    # Check for common validation table headers
                    validation_headers = ['input', 'validation', 'remarks', 'description', 'requirement']
                    header_matches = sum(1 for h in header_texts if any(vh in h for vh in validation_headers))
                    
                    if header_matches >= 2:
                        # For validation tables, use balanced column widths with slightly wider middle column
                        suggested_widths[i] = [1, 1.2, 1]
            except Exception:
                pass
    
    return suggested_widths

def parse_column_widths(width_spec):
    """
    Parse a column width specification string into a dictionary mapping table indices to width lists.
    Format: "0:1,2,1;1:2,1,1,2" means table 0 has columns with relative widths [1,2,1]
    and table 1 has columns with relative widths [2,1,1,2]
    
    Args:
        width_spec: A string specifying column widths
        
    Returns:
        dict: Dictionary mapping table indices to lists of column widths
    """
    if not width_spec:
        return None
        
    result = {}
    
    # Split by semicolon to get each table specification
    table_specs = width_spec.split(';')
    
    for spec in table_specs:
        if not spec or ':' not in spec:
            continue
            
        # Split by colon to get table index and column widths
        parts = spec.split(':')
        if len(parts) != 2:
            continue
            
        try:
            table_idx = int(parts[0].strip())
            widths = [float(w.strip()) for w in parts[1].split(',')]
            
            if widths:
                result[table_idx] = widths
        except ValueError:
            print(f"Warning: Could not parse width specification: {spec}")
            continue
            
    return result if result else None

def apply_document_styles(doc):
    """
    Apply consistent styling to the entire document:
    - Times New Roman font
    - Size 12pt for normal text, 14pt for headings
    - Single line spacing (1.0)
    
    Args:
        doc: The docx document object
    """
    # Process all paragraphs in the document
    for paragraph in doc.paragraphs:
        # Check if it's a heading
        is_heading = paragraph.style.name.startswith('Heading')
        
        # Set paragraph spacing to single (1.0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        
        # Apply font formatting to all runs in the paragraph
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            # Use larger font for headings
            if is_heading:
                run.font.size = Pt(14)
            else:
                run.font.size = Pt(12)
    
    # Apply styles to text in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
    
    # Update document styles
    styles = doc.styles
    for style_name in ['Normal', 'Body Text']:
        if style_name in styles:
            style = styles[style_name]
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
            if hasattr(style, 'paragraph_format'):
                style.paragraph_format.line_spacing = 1.0
                style.paragraph_format.space_before = Pt(0)
                style.paragraph_format.space_after = Pt(0)
    
    # Update heading styles
    for i in range(1, 10):  # Heading 1 through Heading 9
        heading_name = f'Heading {i}'
        if heading_name in styles:
            style = styles[heading_name]
            style.font.name = 'Times New Roman'
            style.font.size = Pt(14)
            if hasattr(style, 'paragraph_format'):
                style.paragraph_format.line_spacing = 1.0
    
    return doc

def fix_table_sizes(doc_path, output_path=None, table_column_widths=None, margin_mm=25.4, reduce_font=False, apply_styles=True):
    """
    Process a Word document, adjusting all tables to fit within A4 page size.
    
    Args:
        doc_path: Path to the input docx file
        output_path: Path to save the modified file (if None, creates a new file)
        table_column_widths: Dictionary mapping table index to list of column width values
                            Example: {0: [1, 2, 1], 1: [3, 1, 1, 1]} sets custom widths for
                            first and second tables in the document
        margin_mm: Page margins in millimeters (default: 25.4mm = 1 inch)
        reduce_font: Whether to reduce font size in table cells to help fit content
        apply_styles: Whether to apply Times New Roman font and spacing settings
    
    Returns:
        str: Path to the saved output file
    """
    if not output_path:
        filename, ext = os.path.splitext(doc_path)
        output_path = f"{filename}_adjusted{ext}"
    
    # Load the document
    try:
        doc = docx.Document(doc_path)
    except Exception as e:
        print(f"Error opening document: {e}")
        return None
    
    # Get A4 dimensions (width and height)
    max_width, max_height = get_a4_dimensions(margin_mm)
    
    # Apply pandoc-specific fixes
    doc = fix_pandoc_tables(doc)
    
    # Apply document styles if requested
    if apply_styles:
        print("Applying Times New Roman, 12pt font with single spacing to document...")
        doc = apply_document_styles(doc)
    
    # If auto-detect is requested but not provided, detect table patterns
    if table_column_widths is None:
        table_column_widths = detect_table_patterns(doc)
        if table_column_widths:
            print(f"Auto-detected table patterns for {len(table_column_widths)} tables")
    
    # Counter for adjusted tables
    adjusted_count = 0
    complex_count = 0
    
    # Collect information about tables for reporting
    table_info = []
    
    # Process all tables in the document
    for i, table in enumerate(doc.tables):
        try:
            # Count max columns in this table
            max_cols = 0
            for row in table.rows:
                max_cols = max(max_cols, len(row.cells))
                
            # Detect variable column structure
            variable_cols = False
            for row in table.rows:
                if len(row.cells) > 0 and len(row.cells) != max_cols:
                    variable_cols = True
                    break
            
            # Get custom column widths for this table if specified
            column_widths = table_column_widths.get(i) if table_column_widths else None
            
            # Check if this is likely a markdown table based on structure
            is_markdown_table = False
            if len(table.rows) > 0 and max_cols > 0:
                # Look for common markdown table patterns
                if len(table.rows) >= 2:
                    # Check cell content for markdown-like formats
                    is_markdown_table = True
            
            # Reduce font size if requested and it's a markdown-converted table
            if reduce_font and (is_markdown_table or variable_cols):
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                # Reduce font size to help fit content
                                current_size = run.font.size.pt if run.font.size else 11
                                run.font.size = Pt(min(current_size, 9))
            
            # Adjust the table with appropriate column widths
            if isinstance(column_widths, str) and column_widths == 'complex-program-spec':
                # Use special handling for complex program spec tables
                handle_complex_table_structure(table, max_width)
                complex_count += 1
            else:
                # Use standard adjustment with specified column widths
                adjust_table_size(table, max_width, column_widths)
            
            adjusted_count += 1
            
            # Record info for reporting
            table_info.append({
                'index': i,
                'max_cols': max_cols,
                'rows': len(table.rows),
                'variable_cols': variable_cols,
                'is_markdown': is_markdown_table
            })
        except Exception as e:
            print(f"Warning: Error adjusting table {i}: {e}")
    
    # Save the modified document
    try:
        doc.save(output_path)
        print(f"Adjusted {adjusted_count} tables in the document ({complex_count} complex tables)")
        
        # Print table info
        print("\nTable Information:")
        for info in table_info:
            print(f"Table {info['index']}: {info['max_cols']} max columns, {info['rows']} rows, " 
                  f"{'Variable columns' if info['variable_cols'] else 'Regular columns'}, " 
                  f"{'Markdown-converted' if info['is_markdown'] else 'Standard'}")
        
        print(f"\nSaved modified document to: {output_path}")
        return output_path
    except Exception as e:
        print(f"Error saving document: {e}")
        return None

def main():
    if len(sys.argv) < 2:
        print("""
Usage: python docxadjust.py <input_docx_file> [output_docx_file] [options]

Arguments:
  input_docx_file   - Path to the input Word document
  output_docx_file  - (Optional) Path to save the modified document
  
Options:
  --widths=SPECS    - Column width specifications
                      Format: "0:1,2,1;1:2,1,1,2" 
                      This means table 0 has columns with relative widths [1,2,1]
                      and table 1 has columns with relative widths [2,1,1,2]
  
  --margin=SIZE     - Page margins in millimeters (default: 25.4mm = 1 inch)
  
  --reduce-font     - Reduce font size in table cells to help fit content
  
  --auto-detect     - Automatically detect table patterns and suggest column widths
                      Works well with pandoc-generated markdown tables
  
  --no-styles       - Do not apply Times New Roman font and spacing settings
                      By default, all text is set to Times New Roman, 12pt (14pt for headings)
                      with single line spacing

Examples:
  python docxadjust.py document.docx
  python docxadjust.py document.docx output.docx
  python docxadjust.py document.docx output.docx --widths="0:1,2,1;1:2,1,1,2"
  python docxadjust.py document.docx --auto-detect --reduce-font
  python docxadjust.py document.docx --margin=20 --reduce-font
        """)
        sys.exit(1)
    
    # Get input file
    input_file = sys.argv[1]
    
    # Get output file (or None for default naming)
    output_file = None
    next_arg_is_output = False
    
    # Parse arguments
    margin_mm = 25.4  # Default: 1 inch margins
    reduce_font = False
    auto_detect = False
    apply_styles = True  # Default: apply Times New Roman and spacing
    column_widths = None
    
    i = 2
    while i < len(sys.argv):
        arg = sys.argv[i]
        
        if next_arg_is_output:
            output_file = arg
            next_arg_is_output = False
        elif arg.startswith('--'):
            # Handle option
            if arg.startswith('--widths='):
                width_spec = arg[len('--widths='):]
                column_widths = parse_column_widths(width_spec)
            elif arg == '--reduce-font':
                reduce_font = True
            elif arg == '--auto-detect':
                auto_detect = True
            elif arg == '--no-styles':
                apply_styles = False
            elif arg.startswith('--margin='):
                try:
                    margin_mm = float(arg[len('--margin='):])
                except ValueError:
                    print(f"Warning: Invalid margin value '{arg[len('--margin='):]}', using default")
            else:
                print(f"Warning: Unknown option '{arg}'")
        elif output_file is None:
            output_file = arg
        
        i += 1
    
    if not os.path.exists(input_file):
        print(f"Error: Input file '{input_file}' not found")
        sys.exit(1)
    
    # Handle auto-detection of table patterns
    if auto_detect:
        print("Auto-detecting table patterns...")
        try:
            doc = docx.Document(input_file)
            detected_widths = detect_table_patterns(doc)
            if detected_widths:
                print(f"Detected column width patterns for {len(detected_widths)} tables:")
                for table_idx, widths in detected_widths.items():
                    if isinstance(widths, str):
                        print(f"  Table {table_idx}: Complex table structure detected")
                    else:
                        width_str = ", ".join(f"{w:.1f}" for w in widths)
                        print(f"  Table {table_idx}: [{width_str}]")
                
                # Merge with any manually specified widths
                if column_widths:
                    for table_idx, widths in detected_widths.items():
                        if table_idx not in column_widths:
                            column_widths[table_idx] = widths
                else:
                    column_widths = detected_widths
            else:
                print("No specific table patterns detected")
        except Exception as e:
            print(f"Warning: Error during pattern detection: {e}")
    
    # Process the document
    print(f"Processing document with margins: {margin_mm}mm")
    result = fix_table_sizes(input_file, output_file, column_widths, margin_mm, reduce_font, apply_styles)
    
    if result:
        print("\nTable adjustment completed successfully!")
    else:
        print("\nFailed to adjust tables")
        sys.exit(1)

if __name__ == "__main__":
    main()
